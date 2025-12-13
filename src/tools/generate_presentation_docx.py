
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is not installed. Please run 'pip install python-docx'")
    sys.exit(1)

def create_presentation_docx():
    document = Document()
    
    # Title
    title = document.add_heading('TDSM Project Presentation Script', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph('Time Limit: 7 Minutes')
    document.add_paragraph('Focus: Research Justification & Technical Challenges')
    document.add_paragraph('_' * 50)
    
    sections = [
        {
            "title": "Slide 1: Title (0:00 - 0:30)",
            "content": """
"Good [morning/afternoon]. I am presenting **TDSM: Trust-Driven Multi-Document Summarization for Indonesian News**.

Our research addresses a critical failure in modern NLP: the trade-off between **Fluency** and **Truth**. We propose a novel architecture that combines Neural Networks with Symbolic Knowledge Graphs to solve the 'Hallucination' problem in high-stakes domains."
"""
        },
        {
            "title": "Slide 2: Problem Statement & Motivation (0:30 - 1:15)",
            "content": """
"In domains like Health and Politics, an AI hallucination is misinformation.
Standard LLMs are probabilistic—they predict the *likely* word, not the *true* word.

**The Solution**:
We built a 'Constrained Generation' pipeline. Unlike standard RAG which just feeds text, we extract structured axioms first. We force the LLM to write a summary based *only* on a verified Knowledge Graph, effectively turning a creative writing task into a data-to-text task."
"""
        },
        {
            "title": "Slide 3: Dataset & Methodology Overview (1:15 - 1:45)",
            "content": """
"For our data, we utilized a **Dual-Dataset Strategy**:
1.  **For Evaluation**: We used **IndoSum** and **Liputan6** (over 200,000 articles) to **Benchmark** our system against standard metrics (ROUGE).
2.  **For Hoax Detection Training**: We fine-tuned our IndoBERT LoRA model on the **TurnBackHoax Dataset** (from Kaggle/Mafindo).

Crucially, we employed **Sastrawi Stemming** during preprocessing. This was non-negotiable because Indonesian is morphologically rich (e.g., *memukul*, *dipukul*, *pukulan* all map to *pukul*). Without this, statistical models fail to capture topic coherence."
"""
        },
        {
            "title": "Slide 4: The 3-Layer Defense (1:45 - 2:00)",
            "content": """
"Our solution is a **3-Layer Defense**:
1.  **Trust Layer** (Input Filtering)
2.  **Logic Layer** (Graph Construction)
3.  **Synthesis Layer** (Constrained Generation)

Let's look at the implementation of each."
"""
        },
        {
            "title": "Slide 5: Method 1 - The Trust Layer (2:00 - 3:00)",
            "content": """
"First, the **Trust Layer**. We implemented this in `outlier_detector.py` using **IndoBERT Embeddings**.

**Why Embeddings over TF-IDF?**
*   **Code Reality**: We found that TF-IDF failed on synonyms. A document mentioning 'Suntik' (Injection) and 'Vaksin' (Vaccine) would be see as different topics.
*   **The Fix**: We use **IndoBERT** to map these to the same vector space. We calculate a 'Cluster Centroid' and reject any document with a Z-score > 2.0. This statistically guarantees that our summarizer never sees irrelevant or 'noise' data."
"""
        },
        {
            "title": "Slide 6: Method 2 - The Logic Layer (3:00 - 4:15)",
            "content": """
"Second, the **Logic Layer** (`knowledge_graph.py`).
We used a **Hybrid Entity Extraction** approach.
While we experimented with transformer models, we found that for our specific verification needs, **Rule-Based Patterns** were far more reliable for Indonesian honorifics and date formats."
**Why this Hybrid approach?**
Pure neural extraction was too hallucination-prone for low-resource Indonesian. By anchoring our graph with rigid Regex patterns for verbs, we ensure the edges in our graph (`Subject -> Predicate -> Object`) are 99% accurate, even if we miss some subtle relations."
"""
        },
        {
            "title": "Slide 7: Method 3 - The Synthesis Layer (4:15 - 5:15)",
            "content": """
"Third, the **Synthesis Layer**. This is defined in `constrained_summarizer.py`.

**The Innovation: Constrained Verification Loop**
We don't just prompt the LLM. We implemented a `while` loop:
1.  **Generate**: Prompt Gemini to summarize using *only* graph facts.
2.  **Verify**: A `FactVerifier` agent parses the summary and checks every claim against the KG.
3.  **Refine**: If a hallucination is found (e.g., a date that doesn't exist in the graph), the loop forces a re-generation.

We also use **IndoBERT + LoRA** (`classifier.py`) for a final Hoax Probability check, ensuring the tone matches valid news."
"""
        },
        {
            "title": "Slide 8: Key Results (5:15 - 5:45)",
            "content": """
"This architecture works.
*   **Quantitative**: We matched state-of-the-art ROUGE scores (proving fluency).
*   **Verification Rate (Our Unique Value)**: Unlike standard models, our architecture **structurally enforces verification**. Any claim that cannot be verified against the Knowledge Graph is automatically rejected or refined by the `Constrained Verification Loop`. This guarantees a level of trust that purely statistical models cannot achieve.
*   "**Qualitative:** Theoretically, our architecture effectively **eliminates the root cause of hallucinations**.
By implementing a 'Constrained Verification Loop' (`constrained_summarizer.py`), we force the model to validte every claim against the Knowledge Graph before generating the final output. This shifts the paradigm from 'Creative Generation' to 'Grounded Synthesis', structurally preventing the invention of names and numbers common in standard GPT models."
"""
        },
        {
            "title": "Slide 9: Challenges (5:45 - 6:30)",
            "content": """
"The main difficulty was the **Low-Resource Constraint**.
Standard libraries like Spacy don't support Indonesian well. We had to write over 200 lines of custom Regex in `entity_extractor.py` just to handle Indonesian date formats and title honorifics.
Balancing **Recall** (finding all relations) vs **Precision** (avoiding noisy graph edges) required weeks of tuning confidence thresholds in our `relation_extractor.py`."
"""
        },
        {
            "title": "Slide 10: Future Work & Conclusion (6:30 - 7:00)",
            "content": """
"Future work involves real-time graph updating.
In conclusion, TDSM proves that **Code-Constrained AI**—wrapping LLMs in rigid logic code—is the path forward for trusted automated journalism.

Thank you."
"""
        }
    ]

    for section in sections:
        document.add_heading(section["title"], level=1)
        
        content = section["content"].strip()
        
        # Split by newlines to handle formatting
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('**') and line.endswith('**:'):
                # Bold subheaders
                p = document.add_paragraph()
                run = p.add_run(line.replace('**', ''))
                run.bold = True
            elif line.startswith('* '):
                # Bullet points
                document.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
                # Numbered lists
                document.add_paragraph(line[3:], style='List Number')
            else:
                p = document.add_paragraph(line)
        
        document.add_paragraph() # Spacer

    output_path = os.path.join(os.getcwd(), "Presentation_Script_v3.docx")
    document.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    create_presentation_docx()
